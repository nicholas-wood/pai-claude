"""
BCG TDA - Friday call one-pager.
Compact, glance-only reference sheet for during the 11:30am screening call.
Two-column where possible, terse spines, not full prose.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION


NAVY = RGBColor(0x1F, 0x49, 0x7D)
GREY = RGBColor(0x66, 0x66, 0x66)


def heading(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = NAVY
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(1)
    return p


def bullet(doc, text, size=9, bold_lead=None):
    p = doc.add_paragraph(style=None)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = Cm(0.4)
    if bold_lead:
        rb = p.add_run(f"{bold_lead}: ")
        rb.bold = True
        rb.font.size = Pt(size)
    r = p.add_run(text)
    r.font.size = Pt(size)
    return p


def main():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.2)
        section.bottom_margin = Cm(1.2)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

    # Title
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = t.add_run("BCG TDA - Friday Call Reference  |  11:30am  |  Glance only, do not read")
    tr.bold = True
    tr.font.size = Pt(12)
    tr.font.color.rgb = NAVY
    t.paragraph_format.space_after = Pt(4)

    # CV walkthrough spine
    heading(doc, "CV WALKTHROUGH (90 sec) - spine, not script")
    bullet(doc, "Started unsure. Drawn to tech at uni. Deliberate call: solid technical foundation FIRST, to lead later.")
    bullet(doc, "7 yrs defence / intel / gov / banking: dev, delivery, architecture, stakeholder side.")
    bullet(doc, "Gen AI: delivered to customers in a bank, then shipped AccessBuddy solo.")
    bullet(doc, "Strategy + policy at national level through surf lifesaving. No formal authority.")
    bullet(doc, "Terrigal workshop: apps vs data confusion. Stepped in, translated, board could decide.")
    bullet(doc, "CLOSER: strong opinions loosely held. Want to be proven wrong. Want to be dumbest in the room.")

    # Three Whys
    heading(doc, "WHY CONSULTING")
    bullet(doc, "Been pointing here for years. Recent engagement = first time it was called 'consulting'.")
    bullet(doc, "Want learning rate across industries, not depth in one company.")
    bullet(doc, "Want higher problem-solving bar. MBB apprenticeship is the fastest way.")

    heading(doc, "WHY BCG")
    bullet(doc, "Most credibly tech-and-data-led MBB (TDA, BCG X, Henderson Institute).")
    bullet(doc, "Intellectual flexibility: less framework-bound than McKinsey, less siloed than Bain.")
    bullet(doc, "Trini and Nora gave a real read on culture and people. Moved it from option to first choice.")

    heading(doc, "WHY TDA")
    bullet(doc, "Where strategy and technical depth meet. Build vs buy, where to place the AI bet.")
    bullet(doc, "AccessBuddy = AI-build proof. LIMSOC = strategy proof. Translation work = my edge.")
    bullet(doc, "Last twelve months make TDA the obvious next step, not a stretch.")

    # Story spines
    heading(doc, "TOP 4 STORY SPINES (STAR + reflection)")
    bullet(doc, "SLSA board workshop, apps vs data. Whiteboard, 3 layers, examples they knew. Board adopted our structures. Lesson: make the tech boring so strategy becomes possible.", bold_lead="M - Tech translation")
    bullet(doc, "Asset Master at Macquarie, regulatory deadline. Separated constraints / preferences / fears. Future-state architecture, staged. Hit deadline. Lesson: separate constraints from preferences.", bold_lead="I - Ambiguity")
    bullet(doc, "Early Macquarie, right substance, wrong delivery. Didn't read the room. Damaged relationships, took months to repair. Lesson: being right isn't enough; manage how it lands.", bold_lead="J - Failure")
    bullet(doc, "SLSA first ever Product Manager. Laughed at first. Reframed from job title to problems it solves. Hired; role now expanding. Lesson: sell the absence, let them own the conclusion.", bold_lead="L - Persuasion")

    # Questions to ask
    heading(doc, "QUESTIONS TO ASK (use 3-4)")
    bullet(doc, "For experienced hires into TDA, what separates the strongest candidates in the process?")
    bullet(doc, "Based on what you've heard today, what should I over-prepare before the case rounds?")
    bullet(doc, "What does the rest of the process look like? Anything different from a standard case format?")
    bullet(doc, "Is Melbourne TDA weighted more to digital strategy or to data/AI delivery? What's growing fastest?")
    bullet(doc, "How does BCG think about levelling for someone with my background?")
    bullet(doc, "Success for a new TDA hire in the first 6-12 months? Most common reasons strong people stumble?")
    bullet(doc, "A project the Melbourne TDA team is excited about right now?")

    # Closer
    heading(doc, "CLOSING QUESTION (always ask)")
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    r = p.add_run('"Is there anything from our conversation today that gives you hesitation about '
                  'progressing me? I\'d rather address it now than have you wonder after the call."')
    r.font.size = Pt(9)
    r.italic = True

    # Footer reminders
    heading(doc, "REMINDERS")
    bullet(doc, "Concise. Land the point, then stop. Don't over-explain when nervous.", size=9)
    bullet(doc, "You've PIVOTED, not pivoting. Speak as if the move is already underway.", size=9)
    bullet(doc, "Treat this like Round 1, not a warm-up. It is the gate.", size=9)

    out = "/Users/nicholaswood/Documents/workspace/bcg-interview-prep/BCG_TDA_Friday_OnePager.docx"
    doc.save(out)
    print(f"Saved: {out}")


if __name__ == "__main__":
    main()
