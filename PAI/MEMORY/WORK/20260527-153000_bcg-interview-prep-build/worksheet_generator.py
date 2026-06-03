"""
BCG Interview Prep - 20-session worksheet generator (2 sessions/day, 10 working days).

Each session combines drills and case work:
- Mental math (10 min)
- Market sizing OR issue tree (10 min)
- Case work block: opener / math / structure / chart / full case / synthesis (20 min)
- Fit story OR reading synthesis (10 min)
- Self-review (2 min)

Sessions alternate block types so AM and PM on the same day are varied.
Output: BCG_TDA_Daily_Worksheets.docx with cover, 20 worksheets, mental math answer key.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import random


# ---------------------------------------------------------------------------
# Session matrix (20 sessions = 10 working days x AM/PM)
# Tier 1: Foundations | Tier 2: Integration | Tier 3: Advanced/pressure
# ---------------------------------------------------------------------------

SESSIONS = [
    # day, session_label, math_tier, block2_type, case_type, block4_type
    (1, "Mon W1 - AM", 1, "sizing", "opener_generic", "fit_cv"),
    (1, "Mon W1 - PM", 1, "tree", "case_math", "synth_reading"),
    (2, "Tue W1 - AM", 1, "tree", "structure", "fit_why_bcg"),
    (2, "Tue W1 - PM", 1, "sizing", "full_profitability", "synth_reading"),
    (3, "Wed W1 - AM", 1, "sizing", "opener_tda", "fit_leadership"),
    (3, "Wed W1 - PM", 2, "tree", "full_market_entry", "synth_reading"),
    (4, "Thu W1 - AM", 2, "tree", "case_math", "fit_failure"),
    (4, "Thu W1 - PM", 2, "sizing", "full_pricing", "synth_reading"),
    (5, "Fri W1 - AM", 2, "sizing", "chart_interp", "fit_conflict"),
    (5, "Fri W1 - PM", 2, "tree", "full_m_and_a", "synth_reading"),
    (6, "Mon W2 - AM", 2, "tree", "opener_tda_build_buy", "fit_persuade"),
    (6, "Mon W2 - PM", 3, "sizing", "full_tda_build_buy", "synth_reading"),
    (7, "Tue W2 - AM", 3, "sizing", "case_math", "fit_tech_translate"),
    (7, "Tue W2 - PM", 3, "tree", "full_tda_data_platform", "synth_reading"),
    (8, "Wed W2 - AM", 3, "tree", "chart_interp_tda", "fit_ethical"),
    (8, "Wed W2 - PM", 3, "sizing", "full_tda_ai_use_cases", "synth_reading"),
    (9, "Thu W2 - AM", 3, "sizing", "synthesis_drill", "fit_tech_decision"),
    (9, "Thu W2 - PM", 3, "tree", "casey_chatbot", "synth_reading"),
    (10, "Fri W2 - AM", 3, "tree", "full_mock_recorded", "fit_mock_full"),
    (10, "Fri W2 - PM", 3, "sizing", "full_mock_recorded_2", "final_reflection"),
]


# ---------------------------------------------------------------------------
# Content banks
# ---------------------------------------------------------------------------

SIZING_PROMPTS = [
    "Estimate the annual market for premium gym memberships in Melbourne CBD.",
    "Estimate the number of flat whites sold in Melbourne CBD on a typical Tuesday.",
    "Estimate the annual TAM for cloud data warehouse services among ASX 200 companies.",
    "Estimate the number of EV public fast-charging bays needed in Victoria by 2030.",
    "Estimate annual revenue of a national chain of 40 fitness studios in Australia.",
    "Estimate the annual spend on enterprise AI tooling (Copilot-class) by the big four Australian banks.",
    "Estimate the addressable market for surf lifesaving club management software across ANZ.",
    "Estimate the size of the Australian B2B SaaS market for HR tech.",
    "Estimate annual cloud infrastructure spend across the Australian federal government.",
    "Estimate the number of data engineers employed across Australian financial services.",
]

TREE_PROMPTS = [
    "An ASX 100 retailer's online conversion rate has fallen 18% YoY. The CDO wants to know why. Build an issue tree.",
    "A major Australian bank's mobile app NPS has dropped from 42 to 28 in 12 months. Build an issue tree.",
    "A health insurer's data warehouse costs have doubled in 18 months while data volume is up only 35%. Build an issue tree.",
    "An Australian SaaS scale-up's gross margin has declined from 78% to 64% over two years. Build an issue tree.",
    "A state government department's cloud migration is 14 months behind schedule. Build an issue tree.",
    "A logistics company's predictive maintenance ML model is generating 60% false positives. Build an issue tree.",
    "A super fund's data platform supports 14 BI tools used by 200 analysts; the CIO wants to rationalise. Build an issue tree.",
    "An energy retailer's customer churn jumped 6 points after launching a new digital self-serve channel. Build an issue tree.",
    "A mining company is deciding whether a $80M investment in an integrated operations data platform will pay back. Build an issue tree.",
    "A FMCG company's analytics function delivers 60+ dashboards but business leaders say they aren't useful. Build an issue tree.",
]

CASE_BLOCKS = {
    "opener_generic": (
        "Case opener drill - first 3 minutes only",
        "Prompt: A national retailer's profits have dropped 25% over two years. The CEO wants to know why and what to do.\n\n"
        "Practise only the opening: (1) play back the objective, (2) ask 1-2 clarifying questions, (3) state your structure, "
        "(4) state your starting hypothesis. Time-box to 3 minutes. Out loud. Do not continue past 3 min."
    ),
    "opener_tda": (
        "Case opener drill (TDA-flavoured) - first 3 minutes only",
        "Prompt: An Australian bank's CDO is being asked by the board whether to consolidate its 4 customer data platforms "
        "into one. He has 6 weeks to recommend. Where do you start?\n\n"
        "Same drill: objective, clarifying, structure, hypothesis. 3 minutes. Out loud."
    ),
    "opener_tda_build_buy": (
        "Case opener drill (TDA build vs buy) - first 3 minutes only",
        "Prompt: A super fund is debating whether to build an internal AI assistant for its call centre or licence a vendor "
        "solution. The CIO wants a recommendation in 4 weeks. How would you frame this?\n\n"
        "Drill: objective, clarifying, structure, hypothesis. 3 minutes. Out loud."
    ),
    "case_math": (
        "Case math drill",
        "1. A retailer has revenue of $1.4B, gross margin of 38%, fixed costs of $320M. What is operating profit?\n\n"
        "2. The same retailer is considering a digital investment costing $45M upfront and $12M/year to operate. "
        "Expected revenue uplift: 3% of current revenue, at 50% incremental margin. What is the payback period?\n\n"
        "3. If the digital investment delivers half the expected uplift, does it still pay back within 5 years?\n\n"
        "Work it out on paper, narrate your thinking out loud, then check your maths twice."
    ),
    "structure": (
        "Case structuring drill",
        "Prompt: An Australian university wants to grow international student revenue by 20% over 3 years without "
        "compromising domestic quality. Build an issue tree to 3 levels deep and state your starting hypothesis. "
        "10 minutes. Then articulate the 'so what' in 2 sentences."
    ),
    "chart_interp": (
        "Chart interpretation drill",
        "Find any chart from an ASX 200 company's most recent annual report or investor presentation. Spend 10 minutes "
        "narrating out loud: (1) what the chart literally shows, (2) the one key insight, (3) the next question you'd ask "
        "the management team, (4) the commercial implication. Write your 4 answers in the space below."
    ),
    "chart_interp_tda": (
        "Chart interpretation drill (TDA-flavoured)",
        "Find a chart from a BCG, McKinsey, or Bain digital/AI insight piece (search 'BCG AI value' or 'McKinsey state of AI'). "
        "Narrate: (1) what it shows, (2) the insight, (3) the next question, (4) the commercial implication. "
        "Bonus: what would a sceptical CFO push back on?"
    ),
    "full_profitability": (
        "Full solo case - Profitability",
        "Prompt: Australia Post's parcel delivery division has seen profit fall 30% over 3 years despite volume growth of 12%. "
        "The CEO wants to know why and what to do.\n\n"
        "Do the full case solo, on paper, out loud. Target: 30-35 minutes. Structure → diagnose → recommend → risks. "
        "Then check against any profitability case in Cosentino or PrepLounge to calibrate."
    ),
    "full_market_entry": (
        "Full solo case - Market Entry",
        "Prompt: An Australian neobank is considering entering New Zealand. They want a recommendation on whether to enter, "
        "and if so, how (organic, acquisition, partnership).\n\n"
        "Full case, solo, on paper, out loud. 30-35 min."
    ),
    "full_pricing": (
        "Full solo case - Pricing",
        "Prompt: An Australian SaaS company that sells inventory management software to mid-market retailers (50-500 stores) "
        "is repricing its product. Current price is $80k/year flat fee. They want to know what to charge and how to structure it.\n\n"
        "Full case, solo, on paper, out loud. 30-35 min."
    ),
    "full_m_and_a": (
        "Full solo case - M&A",
        "Prompt: A large Australian insurer is considering acquiring a fast-growing insurtech for $450M. The insurtech has "
        "$60M revenue, 40% YoY growth, but is not yet profitable. Should they acquire? At what price?\n\n"
        "Full case, solo, on paper, out loud. 30-35 min."
    ),
    "full_tda_build_buy": (
        "Full solo case - TDA: Build vs Buy",
        "Prompt: A super fund is considering whether to build an internal AI assistant for its 800-person call centre, "
        "or licence one of three vendor solutions. Build cost: $12M over 18 months. Vendor licences: $2-4M/year. "
        "The CIO wants a recommendation in 4 weeks. Frame it, work through it, recommend.\n\n"
        "Full case, solo, 30-40 min."
    ),
    "full_tda_data_platform": (
        "Full solo case - TDA: Data Platform Investment",
        "Prompt: An Australian energy retailer wants to know whether a proposed $80M investment in a unified customer data "
        "platform (replacing 6 legacy systems) will pay back over 5 years. Expected benefits: churn reduction, cross-sell, "
        "ops efficiency, regulatory readiness. CFO is sceptical.\n\n"
        "Full case, solo, 30-40 min. Pay particular attention to the value-at-stake calculation."
    ),
    "full_tda_ai_use_cases": (
        "Full solo case - TDA: AI Use Case Prioritisation",
        "Prompt: A large Australian retailer has 22 candidate AI use cases proposed across marketing, supply chain, "
        "customer service, store ops, and merchandising. Total proposed funding is $140M. The CEO has $50M of capacity. "
        "How does she choose? What does the recommendation look like?\n\n"
        "Full case, solo, 30-40 min. Build a clear prioritisation framework."
    ),
    "synthesis_drill": (
        "Case synthesis drill",
        "Take a full case you've done in the past 7 days. Re-do only the synthesis (the closing recommendation). 5 minutes max. "
        "Format: 'My recommendation is X for three reasons: first..., second..., third.... The main risk is Y, which I'd "
        "mitigate by Z. Next step is...'\n\n"
        "Record it on your phone. Watch back. Note 2 things to sharpen."
    ),
    "casey_chatbot": (
        "BCG Casey / digital case practice",
        "Go to bcg.com/careers and complete one BCG Casey interactive case (or any online digital case practice if Casey is "
        "unavailable). 25-30 minutes. Take notes during: where did you slow down, what did the chatbot probe for, what did "
        "you miss the first time. Casey questions appear in TDA first-rounds disproportionately."
    ),
    "full_mock_recorded": (
        "Full mock case (recorded)",
        "Pair with someone (peer, partner, ex-MBB coach). They give you a TDA-flavoured case prompt cold. You run the full case "
        "30-35 min, recorded on Zoom. Afterwards: (1) ask interviewer for 3 specific bits of feedback, "
        "(2) watch the recording yourself, (3) note 3 things to fix.\n\n"
        "If no partner available today, run a cold case from a casebook on yourself with strict 35 min time-box."
    ),
    "full_mock_recorded_2": (
        "Full mock case (recorded) + interview simulation",
        "Second mock today. Different case type from this morning. Same protocol: full case 30-35 min, recorded, three specific "
        "feedback points, three things to fix.\n\n"
        "After case: 10 min of fit interview simulation (interviewer asks 3 behavioural questions, you answer cold)."
    ),
}

FIT_BLOCKS = {
    "fit_cv": (
        "Fit: CV walkthrough drill",
        "Tell your 90-second CV walkthrough. Out loud. Record on phone. Watch back. Note one thing to fix. Do it again."
    ),
    "fit_why_bcg": (
        "Fit: Why BCG / Why TDA / Why now",
        "Three answers, 60 sec each, out loud, recorded. Watch back. Note one thing to fix per answer."
    ),
    "fit_leadership": (
        "Fit: Leadership through ambiguity",
        "Tell your prepared story for 'leading through ambiguity'. STAR plus reflection. 90 sec. Recorded. Watch back."
    ),
    "fit_failure": (
        "Fit: Failure story",
        "Tell your prepared failure story. STAR plus reflection. 90 sec. Recorded. The reflection is what matters - what did "
        "you learn and how has your behaviour changed."
    ),
    "fit_conflict": (
        "Fit: Conflict with peer or senior",
        "Tell your prepared conflict-resolution story. STAR plus reflection. 90 sec. Recorded."
    ),
    "fit_persuade": (
        "Fit: Persuading a senior stakeholder",
        "Tell your prepared influence-without-authority story. STAR plus reflection. 90 sec. Recorded."
    ),
    "fit_tech_translate": (
        "Fit: Translating tech to non-tech executive (highest-leverage TDA story)",
        "Tell your prepared tech-translation story. STAR plus reflection. 90 sec. Recorded. This is the single most "
        "important TDA behavioural - over-prepare it."
    ),
    "fit_ethical": (
        "Fit: Ethical or difficult judgment call",
        "Tell your prepared ethical-judgment story. STAR plus reflection. 90 sec. Recorded."
    ),
    "fit_tech_decision": (
        "Fit: High-stakes technical decision (TDA-specific)",
        "Tell your prepared high-stakes-technical-decision story. STAR plus reflection. 90 sec. Recorded. "
        "Focus on the structure of how you decided, not just the outcome."
    ),
    "fit_mock_full": (
        "Fit: Full behavioural mock",
        "Pair with someone (or self-administer): 5 fit questions cold, 90 sec answer each, no preparation between. Recorded. "
        "Afterwards: watch back, note the weakest answer, plan how to fix it."
    ),
    "synth_reading": (
        "Reading synthesis",
        "Read ONE AFR, FT Lex, or BCG/McKinsey insight article today (5-7 min read). Then synthesise out loud, recorded: "
        "'The argument is X. The author's three supports are A, B, C. The unstated assumption is Y. If I were challenging "
        "this argument, I'd push on Z.' Under 90 seconds."
    ),
    "final_reflection": (
        "Final reflection - end of 2-week sprint",
        "30 min reflection. Write answers to:\n"
        "1. What are the three case types I'm strongest at? Weakest at?\n"
        "2. What are my three biggest behavioural answers? Which one still feels weak?\n"
        "3. What's my mistake-log pattern? Which mistake keeps repeating?\n"
        "4. What's the highest-leverage thing I'd do in Days 15-21 if I had three more weeks?\n"
        "5. How would I describe my readiness to a coach right now: 1-10, with what evidence?"
    ),
}


# ---------------------------------------------------------------------------
# Mental math generator
# ---------------------------------------------------------------------------

def generate_math_problems(session_index, tier):
    rng = random.Random(1000 + session_index)
    problems = []

    for _ in range(5):
        if tier == 1:
            pct = rng.choice([5, 10, 15, 20, 25, 50])
            base = rng.choice([200, 400, 600, 800, 1000])
        elif tier == 2:
            pct = rng.choice([7, 12, 18, 22, 35, 45])
            base = rng.choice([450, 720, 1250, 1800, 2400])
        else:
            pct = rng.choice([3, 8, 14, 23, 37, 62])
            base = rng.choice([1280, 3400, 5750, 8200, 12500])
        problems.append((f"{pct}% of {base:,}", pct * base / 100))

    for _ in range(5):
        if tier == 1:
            a, b = rng.randint(11, 25), rng.randint(11, 25)
        elif tier == 2:
            a, b = rng.randint(15, 99), rng.randint(11, 25)
        else:
            a, b = rng.randint(25, 99), rng.randint(15, 45)
        problems.append((f"{a} x {b}", a * b))

    for _ in range(4):
        if tier == 1:
            a, b = rng.randint(100, 999), rng.choice([5, 10, 20, 25])
        elif tier == 2:
            a, b = rng.randint(500, 9999), rng.choice([8, 12, 16, 24])
        else:
            a, b = rng.randint(1000, 99999), rng.choice([7, 13, 17, 23])
        problems.append((f"{a:,} / {b}", round(a / b, 2)))

    fc = rng.choice([5000000, 12000000, 25000000])
    contribution = rng.choice([15, 25, 40, 75])
    problems.append(
        (f"Breakeven volume: fixed cost ${fc:,}, contribution ${contribution}/unit",
         round(fc / contribution, 0))
    )

    revenue = rng.choice([45, 120, 360, 750, 1500])
    margin_pct = rng.choice([18, 24, 32, 45])
    problems.append(
        (f"Gross profit on ${revenue}M revenue at {margin_pct}% gross margin",
         round(revenue * margin_pct / 100, 1))
    )

    start = rng.choice([100, 250, 500, 1000])
    mult = rng.choice([1.5, 2.0, 2.5, 3.0])
    end = int(start * mult)
    years = rng.choice([3, 5, 7])
    cagr_pct = ((end / start) ** (1 / years) - 1) * 100
    problems.append(
        (f"CAGR from ${start}M to ${end}M over {years} years",
         f"~{round(cagr_pct, 1)}%")
    )

    market_size = rng.choice([2.4, 8.5, 18.0, 45.0])
    share_pct = rng.choice([4, 7, 12, 22])
    problems.append(
        (f"Revenue at {share_pct}% share of a ${market_size}B market",
         f"${round(market_size * share_pct / 100 * 1000, 0)}M")
    )

    investment = rng.choice([15, 40, 120])
    annual_benefit = rng.choice([3, 8, 18, 32])
    problems.append(
        (f"Payback period: ${investment}M investment, ${annual_benefit}M annual benefit",
         f"{round(investment / annual_benefit, 2)} years")
    )

    cost_a = rng.choice([2.4, 5.6, 12.8])
    cost_b = rng.choice([1.8, 4.2, 9.6])
    problems.append(
        (f"Cost change from ${cost_a}M to ${cost_b}M (% change)",
         f"{round((cost_b - cost_a) / cost_a * 100, 1)}%")
    )

    return problems


# ---------------------------------------------------------------------------
# Word doc helpers
# ---------------------------------------------------------------------------

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    return h


def add_section_heading(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    return p


def add_paragraph(doc, text, size=10, italic=False, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.italic = italic
    run.bold = bold
    p.paragraph_format.space_after = Pt(2)
    return p


def add_workspace_lines(doc, count=3):
    for _ in range(count):
        p = doc.add_paragraph("_" * 95)
        for run in p.runs:
            run.font.color.rgb = RGBColor(0xB0, 0xB0, 0xB0)
            run.font.size = Pt(9)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)


def add_math_table(doc, problems):
    table = doc.add_table(rows=10, cols=4)
    table.autofit = True
    table.style = 'Light Grid'
    for i in range(10):
        cell_num_l = table.cell(i, 0)
        cell_prob_l = table.cell(i, 1)
        cell_num_l.text = f"{i + 1}."
        cell_prob_l.text = problems[i][0]
        cell_num_r = table.cell(i, 2)
        cell_prob_r = table.cell(i, 3)
        cell_num_r.text = f"{i + 11}."
        cell_prob_r.text = problems[i + 10][0]
        for cell in [cell_num_l, cell_prob_l, cell_num_r, cell_prob_r]:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)
    for row in table.rows:
        row.cells[0].width = Cm(0.7)
        row.cells[1].width = Cm(8.3)
        row.cells[2].width = Cm(0.7)
        row.cells[3].width = Cm(8.3)


def build_worksheet(doc, session_idx, session_meta, sizing_pool, tree_pool):
    day, label, tier, block2_type, case_type, block4_type = session_meta

    add_heading(doc, f"Day {day} - {label}", level=1)
    p = doc.add_paragraph()
    run = p.add_run(
        f"BCG TDA Prep | Session {session_idx + 1} of 20 | ~50 min | "
        f"Tier {tier}: {'Foundations' if tier == 1 else 'Integration' if tier == 2 else 'Advanced'}"
    )
    run.font.size = Pt(9)
    run.italic = True
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    # 1. Mental math
    add_section_heading(doc, "1. Mental math (10 min)")
    add_paragraph(doc, "Target under 30 sec per problem. Write your working. Out loud for the last 5.", size=9, italic=True)
    problems = generate_math_problems(session_idx, tier)
    add_math_table(doc, problems)

    # 2. Sizing or tree
    if block2_type == "sizing":
        add_section_heading(doc, "2. Market sizing (10 min)")
        prompt = sizing_pool.pop(0)
        add_paragraph(doc, prompt, size=10, italic=True)
        add_paragraph(doc,
                      "Structure: population -> segment -> frequency or usage -> price or value -> sanity check.",
                      size=9)
    else:
        add_section_heading(doc, "2. Issue tree (10 min)")
        prompt = tree_pool.pop(0)
        add_paragraph(doc, prompt, size=10, italic=True)
        add_paragraph(doc,
                      "MECE branches. 3-4 top-level, each with 2-3 sub-branches. Aim for testable hypotheses.",
                      size=9)
    add_workspace_lines(doc, count=3)

    # 3. Case work
    case_title, case_prompt = CASE_BLOCKS[case_type]
    add_section_heading(doc, f"3. Case work (20 min) - {case_title}")
    add_paragraph(doc, case_prompt, size=10, italic=True)
    add_workspace_lines(doc, count=4)

    # 4. Fit or synthesis
    fit_title, fit_prompt = FIT_BLOCKS[block4_type]
    add_section_heading(doc, f"4. Fit or synthesis (10 min) - {fit_title}")
    add_paragraph(doc, fit_prompt, size=10, italic=True)
    add_workspace_lines(doc, count=2)

    # 5. Self-review
    add_section_heading(doc, "5. Self-review (2 min)")
    add_paragraph(doc,
                  "What worked? What broke? One thing to fix next session.",
                  size=9)
    add_workspace_lines(doc, count=2)

    doc.add_page_break()


def build_answer_key(doc):
    add_heading(doc, "Mental Math Answer Key", level=1)
    add_paragraph(
        doc,
        "Answers computed exactly. For percentage and multiplication, accept your answer if within 5% rounding. "
        "For CAGR and business math, within 10% is fine.",
        size=9, italic=True
    )

    for session_idx, session_meta in enumerate(SESSIONS):
        day, label, tier, *_ = session_meta
        problems = generate_math_problems(session_idx, tier)
        add_section_heading(doc, f"Day {day} - {label} (Session {session_idx + 1})")
        table = doc.add_table(rows=10, cols=4)
        table.style = 'Light Grid'
        for i in range(10):
            for col_offset, idx in [(0, i), (2, i + 10)]:
                num_cell = table.cell(i, col_offset)
                ans_cell = table.cell(i, col_offset + 1)
                num_cell.text = f"{idx + 1}."
                ans_cell.text = str(problems[idx][1])
                for cell in [num_cell, ans_cell]:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.font.size = Pt(8)
        for row in table.rows:
            row.cells[0].width = Cm(0.7)
            row.cells[1].width = Cm(8.3)
            row.cells[2].width = Cm(0.7)
            row.cells[3].width = Cm(8.3)
        doc.add_paragraph()


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(1.4)
        section.bottom_margin = Cm(1.4)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)

    # Cover
    title = doc.add_heading('BCG TDA Interview Prep', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run("20 Sessions over 10 Working Days | 2 Sessions/Day | AM + PM")
    sub_run.font.size = Pt(13)
    sub_run.italic = True

    doc.add_paragraph()
    instr = doc.add_paragraph()
    instr.add_run("How to use\n").bold = True
    instr.add_run(
        "- Print this document. Use one worksheet per session.\n"
        "- 10 working days (Mon-Fri x 2 weeks). 2 sessions/day (AM + PM). 10 sessions/week.\n"
        "- Each session ~50 min. Each session combines drills + case work.\n"
        "- Block 1 (Math, 10 min): time yourself, under 30 sec per problem.\n"
        "- Block 2 (Market sizing OR Issue tree, 10 min): work in pen, talk out loud.\n"
        "- Block 3 (Case work, 20 min): rotates across openers, math, structure, charts, full cases, Casey.\n"
        "- Block 4 (Fit OR Synthesis, 10 min): record on phone, watch back, note one fix.\n"
        "- Block 5 (Self-review, 2 min): what worked, what broke, one fix for next session.\n"
        "- Difficulty ramps: Days 1-3 Foundations, Days 4-7 Integration, Days 8-10 Advanced/pressure.\n"
        "- AM and PM on the same day are paired - sizing one session, tree the other; "
        "shorter case work AM, longer case work PM.\n"
        "- Weekends are rest (or light reading only).\n"
        "- Answer key for mental math is at the back."
    )

    doc.add_paragraph()
    sched_h = doc.add_paragraph()
    sched_h.add_run("Schedule at a glance").bold = True
    sched_table = doc.add_table(rows=21, cols=5)
    sched_table.style = 'Light Grid'
    headers = ["Session", "Day", "Time", "Tier", "Case work block"]
    for col, header in enumerate(headers):
        c = sched_table.cell(0, col)
        c.text = header
        for para in c.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(9)
    for idx, session_meta in enumerate(SESSIONS):
        day, label, tier, _, case_type, _ = session_meta
        case_title = CASE_BLOCKS[case_type][0]
        row = sched_table.rows[idx + 1]
        row.cells[0].text = str(idx + 1)
        row.cells[1].text = f"Day {day}"
        row.cells[2].text = label.split(" - ")[-1]
        row.cells[3].text = str(tier)
        row.cells[4].text = case_title
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(8)

    doc.add_page_break()

    # 20 worksheets
    sizing_pool = list(SIZING_PROMPTS)
    tree_pool = list(TREE_PROMPTS)
    for session_idx, session_meta in enumerate(SESSIONS):
        build_worksheet(doc, session_idx, session_meta, sizing_pool, tree_pool)

    # Answer key
    build_answer_key(doc)

    out = "/Users/nicholaswood/Documents/workspace/bcg-interview-prep/BCG_TDA_Daily_Worksheets.docx"
    doc.save(out)
    print(f"Saved: {out}")
    print(f"Sessions: {len(SESSIONS)}")


if __name__ == "__main__":
    main()
